import os
import requests
import string
from bs4 import BeautifulSoup
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

# URLからHTMLを取得
url = "https://jp.quora.com/%E3%83%A2%E3%83%86%E3%82%8B%E7%94%B7%E3%82%92%E3%81%A8%E3%81%93%E3%81%A8%E3%82%93%E7%A0%94%E7%A9%B6%E3%81%97%E3%81%9F%E6%9C%AB%E3%81%AE%E7%B5%90%E8%AB%96%E3%81%AF%E3%81%AA%E3%82%93%E3%81%A7%E3%81%99%E3%81%8B"
response = requests.get(url)
response.encoding = 'utf-8'
soup = BeautifulSoup(response.text, 'html.parser')

# 記事の本文を取得
body_elements = soup.find_all('p')
body = ' '.join([element.text for element in body_elements])
# HTMLからリストを取得
list_elements = soup.find_all(['ul', 'ol'])
lists = [' '.join([li.text for li in element.find_all('li')]) for element in list_elements]
# 記事のタイトルを取得
title = soup.find('h1').text if soup.find('h1') else ''

# スクリプトのあるディレクトリを取得
base_dir = os.path.dirname(os.path.abspath(__file__))

# DataSetディレクトリのパスを作成
dataset = os.path.join(base_dir, 'DataSet')
if not os.path.exists(dataset):
    os.makedirs(dataset)
# タイトルから無効な文字を取り除く
valid_chars = "-_.() %s%s" % (string.ascii_letters, string.digits)
clean_title = ''.join(c for c in title if c in valid_chars)

# タイトルが長すぎる場合は、最初の50文字だけを使用する
max_length = 50
if len(clean_title) > max_length:
    clean_title = clean_title[:max_length]

# リストと本文をテキストファイルに保存
with open(os.path.join(dataset, f'{clean_title}.txt'), 'w', encoding='utf-8') as f:
    f.write(title + '\n' + body)
    for list_item in lists:
        f.write(list_item + '\n')

# docxファイルに保存
doc = Document()
doc.add_heading(title, 0)
doc.add_paragraph(body)
for list_item in lists:
    doc.add_paragraph(list_item)  # リストを段落として追加
doc.save(os.path.join(dataset, f'{clean_title}.docx'))

# フォントファイルのフルパスを指定
font_path = '/Users/kamehaku/Desktop/Add-LLM vs Chat-GPT/tools/Noto_Sans_JP/static/NotoSansJP-Regular.ttf'


# フォントの登録
pdfmetrics.registerFont(TTFont('NotoSansJP', font_path))

# PDFファイルの作成
pdf_file = os.path.join(dataset, f"{clean_title}.pdf")
doc = SimpleDocTemplate(pdf_file, pagesize=letter)

## スタイルの設定
styles = getSampleStyleSheet()
notosans_style = styles['Normal'].clone('NotoSansJP')  # 既存のスタイルをコピー
notosans_style.fontName = 'NotoSansJP'
notosans_style.fontSize = 10
styles.add(notosans_style)  # 新しいスタイルを追加

# タイトル用のスタイルを作成
title_style = styles['Heading1'].clone('NotoSansJP-Title')  # 既存のスタイルをコピー
title_style.fontName = 'NotoSansJP'
title_style.fontSize = 20  # フォントサイズを大きくする
styles.add(title_style)  # 新しいスタイルを追加

# テキストをParagraphオブジェクトに変換
elements = []
elements.append(Paragraph(title, styles['NotoSansJP-Title']))  # タイトル用のスタイルを使用

# 本文とリストの内容を追加
lines = (body + '\n' + '\n'.join(lists)).split('\n')
for line in lines:
    elements.append(Paragraph(line, styles['NotoSansJP']))

# PDFを生成
doc.build(elements)