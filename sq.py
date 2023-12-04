import os
import requests
import string
from bs4 import BeautifulSoup
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
# URLからHTMLを取得
url = "https://ja.wikipedia.org/wiki/2023%E5%B9%B4%E3%81%AE%E6%B0%97%E8%B1%A1%E3%83%BB%E5%9C%B0%E8%B1%A1%E3%83%BB%E5%A4%A9%E8%B1%A1"
response = requests.get(url)
response.encoding = 'utf-8'
soup = BeautifulSoup(response.text, 'html.parser')

# 記事の本文を取得
body_elements = soup.find_all('p')
body = ' '.join([element.text for element in body_elements])
# HTMLからリストを取得
list_elements = soup.find_all(['ul', 'ol'])
lists = [' '.join([li.text for li in element.find_all('li')]) for element in list_elements]
# 記事のタイトルを取得
title = soup.find('h1').text if soup.find('h1') else ''

# スクリプトのあるディレクトリを取得
base_dir = os.path.dirname(os.path.abspath(__file__))

# DataSetディレクトリのパスを作成
dataset = os.path.join(base_dir, 'DataSet')
if not os.path.exists(dataset):
    os.makedirs(dataset)
# タイトルから無効な文字を取り除く
valid_chars = "-_.() %s%s" % (string.ascii_letters, string.digits)
clean_title = ''.join(c for c in title if c in valid_chars)

# リストと本文をテキストファイルに保存
with open(os.path.join(dataset, f'{clean_title}_full.txt'), 'w', encoding='utf-8') as f:
    f.write(title + '\n' + body)
    for list_item in lists:
        f.write(list_item + '\n')

# docxファイルに保存
doc = Document()
doc.add_heading(title, 0)
doc.add_paragraph(body)
for list_item in lists:
    doc.add_paragraph(list_item)  # リストを段落として追加
doc.save(os.path.join(dataset, f'{clean_title}.docx'))

# フォントファイルのフルパスを指定
font_path = '/Users/kamehaku/Desktop/Add-LLM vs Chat-GPT/tools/Noto_Sans_JP/static/NotoSansJP-Regular.ttf'


# 日本語フォントの登録
pdfmetrics.registerFont(TTFont('NotoSansJP', font_path))


# pdfファイルに保存
c = canvas.Canvas(os.path.join(dataset, f"{clean_title}.pdf"), pagesize=letter)
width, height = letter
c.setFont("NotoSansJP", 12)  # 日本語フォントを使用
c.drawString(30, height - 50, title)
textobject = c.beginText()
textobject.setTextOrigin(30, height - 70)
textobject.setFont("NotoSansJP", 10)  # 日本語フォントを使用

# テキストを行単位で処理し、必要に応じて新しいページを開始
lines = (body + '\n' + '\n'.join(lists)).split('\n')
current_height = height - 70
for line in lines:
    # ページの下端に達したら新しいページを開始
    if current_height < 40:
        c.drawText(textobject)
        c.showPage()
        c.setFont("NotoSansJP", 12)
        textobject = c.beginText()
        textobject.setTextOrigin(30, height - 70)
        textobject.setFont("NotoSansJP", 10)
        current_height = height - 70
    textobject.textLine(line)
    current_height -= 12

# 最後のページのテキストを描画
c.drawText(textobject)
c.save()